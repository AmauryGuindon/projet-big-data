"""Génère le rapport technique Big Data au format Word (.docx)."""
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Inches

ROOT = Path(__file__).parent
DIAG = ROOT / "diagrams"
OUTPUT = ROOT / "Rapport_Big_Data.docx"

# Couleurs cohérentes avec les diagrammes
COL_PRIMARY = RGBColor(0x2C, 0x3E, 0x50)   # bleu foncé sobre
COL_ACCENT = RGBColor(0xE6, 0x73, 0x22)    # orange (titre couverture)
COL_SUBTITLE = RGBColor(0x4A, 0x90, 0xE2)  # bleu
COL_TEXT = RGBColor(0x1A, 0x1A, 0x1A)
COL_MUTED = RGBColor(0x66, 0x66, 0x66)
COL_CODE_BG = "F4F4F4"

doc = Document()

# ---------- Styles globaux ----------
styles = doc.styles

# Police par défaut
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = COL_TEXT
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

# Titre 1
h1 = styles["Heading 1"]
h1.font.name = "Calibri"
h1.font.size = Pt(20)
h1.font.bold = True
h1.font.color.rgb = COL_PRIMARY
h1.paragraph_format.space_before = Pt(18)
h1.paragraph_format.space_after = Pt(10)
h1.paragraph_format.keep_with_next = True

# Titre 2
h2 = styles["Heading 2"]
h2.font.name = "Calibri"
h2.font.size = Pt(14)
h2.font.bold = True
h2.font.color.rgb = COL_SUBTITLE
h2.paragraph_format.space_before = Pt(12)
h2.paragraph_format.space_after = Pt(6)

# Titre 3
h3 = styles["Heading 3"]
h3.font.name = "Calibri"
h3.font.size = Pt(12)
h3.font.bold = True
h3.font.italic = False
h3.font.color.rgb = COL_PRIMARY

# Marges
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)


# ---------- Helpers ----------
def add_page_break():
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_para(text, bold=False, italic=False, size=None, color=None,
             align=None, space_after=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_bullets(items, style="List Bullet"):
    for it in items:
        p = doc.add_paragraph(style=style)
        if isinstance(it, tuple):
            label, rest = it
            r1 = p.add_run(label)
            r1.bold = True
            p.add_run(" — " + rest)
        else:
            p.add_run(it)


def add_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = COL_MUTED


def add_image(filename, width_cm=15.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(DIAG / filename), width=Cm(width_cm))


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_table(headers, rows, header_bg="2C3E50", header_color=RGBColor(0xFF, 0xFF, 0xFF),
              col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # En-têtes
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = header_color
        run.font.size = Pt(10.5)
        set_cell_bg(cell, header_bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Lignes
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Largeur colonnes
    if col_widths:
        for col_idx, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[col_idx].width = Cm(w)
    return table


def add_code_block(code, language=""):
    """Bloc de code monospace, fond gris clair, encadré."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.4)

    # Fond gris via shading sur le paragraphe
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), COL_CODE_BG)
    p_pr.append(shd)

    # Bordure
    pbdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), "CCCCCC")
        pbdr.append(b)
    p_pr.append(pbdr)

    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def add_toc():
    """Insère un champ TOC. L'utilisateur fait F9 dans Word pour le mettre à jour."""
    para = doc.add_paragraph()
    run = para.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'

    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "Faites F9 (ou clic droit → Mettre à jour les champs) pour générer le sommaire."

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_sep)
    run._r.append(placeholder)
    run._r.append(fld_char_end)


# ===================================================================
# COUVERTURE
# ===================================================================
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Projet Big Data")
r.font.size = Pt(16)
r.font.color.rgb = COL_MUTED
r.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
r = p.add_run("Rapport technique")
r.font.size = Pt(44)
r.font.bold = True
r.font.color.rgb = COL_ACCENT

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
r = p.add_run("Plateforme distribuée d'ingestion et d'analyse de données massives")
r.font.size = Pt(14)
r.font.color.rgb = COL_PRIMARY
r.italic = True

for _ in range(8):
    doc.add_paragraph()

for label in ["GUINDON Amaury", "Promotion M1 Big Data IA",
              "École Supérieure du Numérique 89", "Année 2025-2026"]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label)
    r.font.size = Pt(14)
    r.font.color.rgb = COL_SUBTITLE

add_page_break()

# ===================================================================
# SOMMAIRE
# ===================================================================
p = doc.add_paragraph()
r = p.add_run("Sommaire")
r.font.size = Pt(24)
r.bold = True
r.font.color.rgb = COL_PRIMARY
p.paragraph_format.space_after = Pt(18)

add_toc()
add_page_break()

# ===================================================================
# 1. CONTEXTE ET OBJECTIFS
# ===================================================================
doc.add_heading("1. Présentation du contexte et des objectifs", level=1)

doc.add_heading("1.1 Contexte", level=2)
add_para(
    "Les entreprises modernes collectent et exploitent des volumes de données sans cesse "
    "croissants : transactions, événements applicatifs, flux IoT, catalogues produits, contenus "
    "multimédia. Les systèmes de gestion de données relationnels traditionnels ne suffisent plus "
    "à absorber ces volumes ni à garantir simultanément les exigences de haute disponibilité, "
    "de scalabilité horizontale et de tolérance aux pannes attendues d'une plateforme moderne."
)
add_para(
    "Ce projet, réalisé dans le cadre du module « Big Data » du Master 1 Big Data IA, vise à "
    "concevoir et implémenter une plateforme distribuée capable de couvrir l'ensemble du cycle "
    "de vie d'une donnée : ingestion, stockage scalable, archivage, et requêtes analytiques. "
    "L'architecture mobilise trois briques fondamentales du paysage Big Data actuel : une base "
    "distribuée NoSQL (Apache Cassandra), un stockage objet de type Data Lake (MinIO / "
    "compatible Amazon S3), et une couche de conteneurisation reproductible (Docker)."
)

doc.add_heading("1.2 Problématique", level=2)
add_para(
    "Une entreprise fictive — opérant à la fois dans la vente en ligne et la distribution de "
    "produits de mode — doit mettre en place une plateforme répondant aux besoins suivants :"
)
add_bullets([
    "collecter des données de natures hétérogènes (transactions structurées, catalogues semi-structurés, contenus multimédia non structurés) ;",
    "stocker ces données de manière scalable pour absorber leur croissance dans le temps ;",
    "permettre des requêtes analytiques rapides sur des dimensions métier (pays, marque, catégorie) ;",
    "garantir la disponibilité du service même en cas de panne d'un nœud ;",
    "être déployable de manière reproductible sur n'importe quel environnement.",
])

doc.add_heading("1.3 Objectifs du projet", level=2)
add_para("Concrètement, le projet doit livrer :")
add_bullets([
    ("une architecture cohérente",
     "intégrant ingestion, stockage distribué et stockage objet ;"),
    ("un modèle de données Cassandra",
     "orienté requêtes, démontrant la maîtrise des concepts de partition et de clustering ;"),
    ("un dépôt Docker reproductible",
     "permettant à un évaluateur de relancer l'ensemble en une commande ;"),
    ("une argumentation Big Data",
     "ancrée dans les principes des 5V, du théorème CAP et de la scalabilité horizontale ;"),
    ("un jeu de requêtes CQL réelles",
     "exécutables sur les données effectivement ingérées."),
])

add_page_break()

# ===================================================================
# 2. ARCHITECTURE
# ===================================================================
doc.add_heading("2. Description détaillée de l'architecture", level=1)

doc.add_heading("2.1 Vue d'ensemble", level=2)
add_para(
    "La plateforme se compose de quatre éléments principaux, tous conteneurisés et orchestrés "
    "via docker-compose. Les sources de données brutes sont consommées par un service "
    "d'ingestion écrit en Python ; celui-ci écrit simultanément les enregistrements structurés "
    "dans Cassandra (pour les requêtes analytiques) et archive les fichiers d'origine dans le "
    "stockage objet MinIO (pour le rejouage, l'archivage légal et la séparation stockage/calcul)."
)
add_image("01_architecture.png", width_cm=16.0)
add_caption("Figure 1 — Vue globale de l'architecture")

doc.add_heading("2.2 Composants", level=2)
add_bullets([
    ("Sources",
     "trois jeux de données réalistes représentant les trois familles de données rencontrées en "
     "production : un CSV de transactions e-commerce (structuré), un dump JSON du catalogue "
     "Flipkart Fashion (semi-structuré), et un corpus d'images JPG accompagnées de leurs "
     "métadonnées CSV (non structuré)."),
    ("Service d'ingestion (Python 3.11)",
     "construit dans une image Docker dédiée. Il lit les sources, applique un parsing et un "
     "nettoyage minimal (gestion des valeurs nulles, conversion des types), puis insère les "
     "données dans Cassandra par lots de 100 enregistrements (BatchStatement), tout en "
     "téléversant les fichiers bruts vers MinIO."),
    ("Apache Cassandra 4.1",
     "base distribuée NoSQL choisie pour sa scalabilité linéaire et son modèle AP (Availability "
     "+ Partition tolerance). Une étape d'initialisation (service cassandra-init) crée le "
     "keyspace et les tables au premier démarrage."),
    ("MinIO",
     "stockage objet compatible API Amazon S3, exécuté localement. Le pipeline est prévu pour "
     "basculer vers AWS S3 réel simplement en redéfinissant S3_ENDPOINT_URL et les "
     "identifiants AWS."),
    ("Volumes Docker",
     "deux volumes persistants (cassandra_data et minio_data) garantissent la persistance des "
     "données au-delà du cycle de vie des conteneurs."),
])

doc.add_heading("2.3 Flux de données", level=2)
add_para(
    "Le pipeline est implémenté sous la forme d'un job one-shot : il s'exécute, charge les "
    "données, puis s'arrête (`restart: \"no\"`). Cette stratégie convient à un MVP et peut être "
    "remplacée par un consommateur de stream (Kafka, Pulsar) lors d'un passage en production."
)
add_image("02_flux_donnees.png", width_cm=16.0)
add_caption("Figure 2 — Flux de données du service d'ingestion")

add_para(
    "Pour chaque type de donnée, le flux est le suivant : lecture en streaming du fichier "
    "source, conversion typée des champs, validation minimale (rejet des enregistrements sans "
    "clé), regroupement en batch, insertion Cassandra et — en parallèle — copie du fichier "
    "brut vers le bucket S3 (« raw zone »). Les images sont en plus téléversées individuellement "
    "et leur clé S3 est stockée dans la table de métadonnées Cassandra, ce qui matérialise un "
    "pattern classique de lien entre métadonnées chaudes (Cassandra) et binaires froids (Data Lake)."
)

add_page_break()

# ===================================================================
# 3. JUSTIFICATION DES CHOIX TECHNIQUES
# ===================================================================
doc.add_heading("3. Justification des choix techniques", level=1)

doc.add_heading("3.1 Pourquoi Apache Cassandra ?", level=2)
add_para(
    "Cassandra est une base NoSQL orientée colonnes wide-column, conçue dès l'origine pour "
    "une distribution multi-nœuds. Elle a été retenue ici pour quatre raisons principales :"
)
add_bullets([
    ("Scalabilité horizontale linéaire",
     "ajouter un nœud augmente proportionnellement la capacité de stockage et le débit, sans "
     "coordinateur central ni single point of failure."),
    ("Disponibilité prioritaire (modèle AP)",
     "Cassandra reste accessible en écriture même en cas de partition réseau, ce qui correspond "
     "au besoin métier d'un service de collecte continu."),
    ("Modélisation orientée requêtes",
     "le modèle physique est dérivé des requêtes à servir, ce qui garantit des temps de réponse "
     "constants quelle que soit la taille du dataset, à la différence des SGBDR où la latence "
     "croît avec le volume."),
    ("Réplication native et configurable",
     "le facteur de réplication et la stratégie (SimpleStrategy, NetworkTopologyStrategy) sont "
     "définis au niveau du keyspace, sans coût applicatif."),
])

doc.add_heading("3.2 Pourquoi MinIO (S3-compatible) ?", level=2)
add_para(
    "Le besoin d'un stockage objet est lié à trois principes structurants d'une plateforme Big "
    "Data : la séparation stockage / calcul, la conservation des données brutes en vue d'un "
    "rejouage ou d'un retraitement futur, et la prise en charge native des fichiers de grande "
    "taille (images, archives) qui ne sont pas adaptés à un stockage relationnel ou NoSQL "
    "transactionnel."
)
add_para(
    "MinIO a été préféré à AWS S3 pour la phase de développement pour deux raisons : aucun "
    "coût ni compte cloud à provisionner, et une API strictement compatible S3 qui permet de "
    "basculer vers AWS en production sans modification du code (seule la variable "
    "S3_ENDPOINT_URL est à neutraliser)."
)

doc.add_heading("3.3 Pourquoi Docker et docker-compose ?", level=2)
add_para(
    "La conteneurisation répond à l'exigence de reproductibilité posée par le cahier des "
    "charges. Chaque composant (Cassandra, init Cassandra, MinIO, ingestor) est isolé dans son "
    "conteneur avec ses dépendances figées. docker-compose orchestre le démarrage ordonné — "
    "le service ingestor n'est lancé qu'une fois Cassandra prête (healthcheck) et que le schéma "
    "est appliqué (depends_on … condition: service_completed_successfully). Le tout est "
    "déployable sur une machine standard via une seule commande."
)

doc.add_heading("3.4 Pourquoi Python pour l'ingestion ?", level=2)
add_para(
    "Python a été retenu pour la couche d'ingestion en raison de son écosystème mature pour le "
    "Big Data (drivers Cassandra et S3 officiels, manipulation aisée des formats CSV / JSON), "
    "de sa lisibilité, et du coût d'entrée minimal en équipe. Pour des volumes "
    "significativement plus élevés, le service pourrait être remplacé par un job Spark ou un "
    "consommateur Kafka Streams sans modifier le modèle Cassandra."
)

add_page_break()

# ===================================================================
# 4. ANALYSE BIG DATA (5V, CAP)
# ===================================================================
doc.add_heading("4. Analyse au regard des principes Big Data", level=1)

doc.add_heading("4.1 Les 5V appliqués au projet", level=2)
add_table(
    headers=["Dimension", "Application au projet"],
    rows=[
        ("Volume",
         "Le pipeline est dimensionné pour absorber des millions d'enregistrements. Les limites "
         "actuelles (CSV_LIMIT, JSON_LIMIT, IMAGE_LIMIT) sont des garde-fous de démonstration "
         "désactivables en passant la valeur à -1. Cassandra et S3 sont nativement conçus pour "
         "le téraoctet et au-delà."),
        ("Vélocité",
         "L'ingestion s'effectue en mode batch avec BatchStatement par paquets de 100 lignes, "
         "stratégie qui amortit le coût réseau. Une évolution vers un consommateur Kafka "
         "permettrait de passer du batch au quasi-temps réel sans modifier le modèle de données."),
        ("Variété",
         "Trois formats représentatifs sont couverts : structuré (CSV transactions), "
         "semi-structuré (JSON catalogue), non structuré (images JPG + métadonnées). Chaque "
         "format est ingéré par un pipeline dédié et stocké de la manière la plus adaptée."),
        ("Véracité",
         "Le pipeline applique un parsing défensif (rejet des lignes sans clé primaire, "
         "tolérance aux valeurs manquantes, normalisation des chaînes). La donnée brute est "
         "conservée en parallèle sur S3, ce qui permet un retraitement complet en cas de bug du "
         "pipeline."),
        ("Valeur",
         "Le modèle Cassandra est conçu pour servir des requêtes métier à forte valeur : "
         "ventes par pays et par mois, disponibilité produit par marque, catalogue image par "
         "catégorie et genre. Ces requêtes sont directement actionnables par un analyste."),
    ],
    col_widths=[3.0, 13.0],
)

doc.add_heading("4.2 Théorème CAP", level=2)
add_para(
    "Le théorème CAP énonce qu'un système distribué ne peut garantir simultanément que deux des "
    "trois propriétés suivantes : Consistency (cohérence), Availability (disponibilité), "
    "Partition tolerance (tolérance au partitionnement). En présence inévitable de partitions "
    "réseau dans un système distribué, le choix se résume à un arbitrage entre C et A."
)
add_image("04_cap_theorem.png", width_cm=11.0)
add_caption("Figure 3 — Positionnement de Cassandra dans le théorème CAP")

add_para(
    "Cassandra se positionne comme un système AP : elle privilégie la disponibilité et la "
    "tolérance au partitionnement, au prix d'une cohérence dite éventuelle (eventual "
    "consistency). Concrètement, deux nœuds peuvent temporairement présenter des versions "
    "différentes d'une même ligne, l'alignement se faisant via les mécanismes d'anti-entropie "
    "(read repair, hinted handoff). Ce compromis est cohérent avec le cas d'usage : il est "
    "préférable que l'analyste reçoive une donnée légèrement décalée plutôt que de voir le "
    "service tomber lors d'un incident réseau."
)
add_para(
    "Pour les rares cas nécessitant une cohérence plus forte, Cassandra permet de spécifier au "
    "niveau de chaque requête un niveau de cohérence (ONE, QUORUM, ALL), ce qui offre un "
    "curseur applicatif entre les deux extrêmes."
)

doc.add_heading("4.3 Scalabilité", level=2)
add_para(
    "La scalabilité est obtenue à plusieurs niveaux. Sur Cassandra, l'ajout d'un nœud "
    "redistribue automatiquement les partitions via le placement par hachage (Murmur3) sur un "
    "anneau de tokens. Sur MinIO, le mode distribué permet d'agréger plusieurs serveurs "
    "physiques en un cluster unique. Sur l'ingestion enfin, le découpage par type de données "
    "permet de paralléliser horizontalement (un conteneur ingestor par flux). Aucune brique "
    "n'introduit de single point of failure structurel."
)

add_page_break()

# ===================================================================
# 5. MODÉLISATION CASSANDRA
# ===================================================================
doc.add_heading("5. Modélisation Cassandra", level=1)

doc.add_heading("5.1 Principes appliqués", level=2)
add_para(
    "À la différence d'une modélisation relationnelle qui part des entités métier, la "
    "modélisation Cassandra part des requêtes à servir. Chaque requête principale donne lieu à "
    "une table dédiée, dénormalisée si nécessaire. Trois règles d'or ont guidé la conception :"
)
add_bullets([
    ("Partition key bien choisie",
     "la clé de partition détermine sur quel nœud les données seront stockées. Elle doit être "
     "suffisamment sélective pour répartir la charge (éviter les hotspots) et corrélée à la "
     "requête (toute requête doit pouvoir cibler une seule partition)."),
    ("Taille de partition maîtrisée",
     "une partition Cassandra ne devrait pas dépasser environ 100 Mo ni 100 000 lignes. Les "
     "partition keys composites du projet (pays + mois, marque + disponibilité, catégorie + "
     "genre) maintiennent les partitions à une taille raisonnable."),
    ("Clustering keys ordonnées",
     "les colonnes de clustering définissent l'ordre physique au sein d'une partition, ce qui "
     "permet des range queries efficaces (par exemple, ventes d'un pays sur un mois triées par "
     "date décroissante)."),
])

doc.add_heading("5.2 Schéma physique", level=2)
add_image("03_modele_cassandra.png", width_cm=16.0)
add_caption("Figure 4 — Schéma physique des trois tables Cassandra")

doc.add_heading("5.3 Détail des tables", level=2)

doc.add_heading("Table 1 — sales_by_country_month", level=3)
add_para(
    "Cette table porte les transactions e-commerce. La clé de partition composite "
    "(country, year_month) co-localise toutes les ventes d'un même pays pour un mois donné, ce "
    "qui correspond exactement à la requête analytique principale. Les clustering keys "
    "(invoice_date DESC, invoice_no, stock_code) donnent un ordre chronologique inversé dans "
    "la partition et garantissent l'unicité de chaque ligne de facture."
)
add_code_block(
    "CREATE TABLE sales_by_country_month (\n"
    "  country text,\n"
    "  year_month text,\n"
    "  invoice_date timestamp,\n"
    "  invoice_no text,\n"
    "  stock_code text,\n"
    "  description text,\n"
    "  quantity int,\n"
    "  unit_price decimal,\n"
    "  customer_id text,\n"
    "  total_amount decimal,\n"
    "  PRIMARY KEY ((country, year_month), invoice_date, invoice_no, stock_code)\n"
    ") WITH CLUSTERING ORDER BY (invoice_date DESC, invoice_no ASC, stock_code ASC);"
)

doc.add_heading("Table 2 — products_by_brand_availability", level=3)
add_para(
    "Cette table sert les requêtes catalogue : « pour la marque X, lister les produits "
    "(in)disponibles ». La clé de partition (brand, out_of_stock) sépare physiquement les "
    "produits disponibles des indisponibles, ce qui rend la requête de stock à un seul accès "
    "disque. La clustering key pid garantit l'unicité par produit dans la partition."
)
add_code_block(
    "CREATE TABLE products_by_brand_availability (\n"
    "  brand text,\n"
    "  out_of_stock boolean,\n"
    "  pid text,\n"
    "  title text,\n"
    "  selling_price int,\n"
    "  actual_price int,\n"
    "  average_rating decimal,\n"
    "  category text,\n"
    "  sub_category text,\n"
    "  crawled_at timestamp,\n"
    "  PRIMARY KEY ((brand, out_of_stock), pid)\n"
    ");"
)

doc.add_heading("Table 3 — image_metadata_by_category_gender", level=3)
add_para(
    "Cette table joue le rôle de catalogue d'images : elle indexe par (category, gender) les "
    "métadonnées de chaque image et stocke la clé S3 (image_object_key) qui pointe vers le "
    "binaire dans le Data Lake. Cette séparation des métadonnées chaudes (Cassandra) et des "
    "binaires froids (S3) est un pattern standard dans les architectures de type Data Lake."
)
add_code_block(
    "CREATE TABLE image_metadata_by_category_gender (\n"
    "  category text,\n"
    "  gender text,\n"
    "  product_id int,\n"
    "  product_type text,\n"
    "  colour text,\n"
    "  usage text,\n"
    "  product_title text,\n"
    "  image_name text,\n"
    "  image_object_key text,\n"
    "  image_url text,\n"
    "  PRIMARY KEY ((category, gender), product_id)\n"
    ");"
)

doc.add_heading("5.4 Stratégie de réplication", level=2)
add_para(
    "Le keyspace est actuellement créé avec une SimpleStrategy et un facteur de réplication "
    "de 1, ce qui est suffisant pour un développement mono-nœud reproductible sur une machine "
    "standard. En production, il conviendrait de basculer sur une NetworkTopologyStrategy avec "
    "un facteur ≥ 3 par centre de données, afin de garantir la tolérance à la perte d'un nœud "
    "et la cohérence QUORUM."
)
add_code_block(
    "-- Configuration actuelle (développement)\n"
    "CREATE KEYSPACE bigdata_project\n"
    "WITH replication = {'class': 'SimpleStrategy', 'replication_factor': 1};\n\n"
    "-- Configuration cible (production multi-DC)\n"
    "ALTER KEYSPACE bigdata_project\n"
    "WITH replication = {\n"
    "  'class': 'NetworkTopologyStrategy',\n"
    "  'dc_eu_west': 3,\n"
    "  'dc_us_east': 3\n"
    "};"
)

add_page_break()

# ===================================================================
# 6. REQUÊTES CQL
# ===================================================================
doc.add_heading("6. Exemples de requêtes CQL", level=1)
add_para(
    "Les requêtes ci-dessous sont disponibles dans le fichier cassandra/queries.cql du dépôt et "
    "exécutables directement via cqlsh dans le conteneur Cassandra. Chaque requête cible une "
    "seule partition, ce qui garantit une latence faible et constante."
)

doc.add_heading("6.1 Ventes d'un pays sur un mois donné", level=2)
add_code_block(
    "SELECT invoice_date, invoice_no, stock_code, quantity, unit_price, total_amount\n"
    "FROM sales_by_country_month\n"
    "WHERE country = 'United Kingdom' AND year_month = '2010-12'\n"
    "LIMIT 20;"
)
add_para(
    "Cette requête exploite directement la partition key et retourne les ventes triées par "
    "date décroissante grâce au CLUSTERING ORDER. Le coût est O(1) en nombre de partitions "
    "interrogées.",
    italic=True, size=10, color=COL_MUTED,
)

doc.add_heading("6.2 Produits indisponibles d'une marque", level=2)
add_code_block(
    "SELECT pid, title, selling_price, average_rating, category, sub_category\n"
    "FROM products_by_brand_availability\n"
    "WHERE brand = 'York' AND out_of_stock = true\n"
    "LIMIT 20;"
)
add_para(
    "La clé composite (brand, out_of_stock) permet d'isoler instantanément les ruptures d'une "
    "marque sans scan, opération typiquement utilisée par un service de réapprovisionnement.",
    italic=True, size=10, color=COL_MUTED,
)

doc.add_heading("6.3 Images d'une catégorie pour un genre", level=2)
add_code_block(
    "SELECT product_id, product_title, image_name, image_object_key, image_url\n"
    "FROM image_metadata_by_category_gender\n"
    "WHERE category = 'Apparel' AND gender = 'Girls'\n"
    "LIMIT 20;"
)
add_para(
    "Cette requête retourne les métadonnées ; l'image binaire elle-même est ensuite récupérée "
    "depuis MinIO à partir de la clé image_object_key. Ce double accès est typique d'une "
    "architecture séparant les métadonnées chaudes des binaires froids.",
    italic=True, size=10, color=COL_MUTED,
)

doc.add_heading("6.4 Comptage par partition (vérification post-ingestion)", level=2)
add_code_block(
    "SELECT COUNT(*) FROM sales_by_country_month;\n"
    "SELECT COUNT(*) FROM products_by_brand_availability;\n"
    "SELECT COUNT(*) FROM image_metadata_by_category_gender;"
)
add_para(
    "Ces requêtes de contrôle sont coûteuses sur un cluster multi-nœuds car elles balayent "
    "l'ensemble des partitions ; elles ne doivent pas être utilisées en production mais "
    "constituent un bon outil de validation pour le présent projet.",
    italic=True, size=10, color=COL_MUTED,
)

add_page_break()

# ===================================================================
# 7. PERFORMANCES & LIMITES
# ===================================================================
doc.add_heading("7. Analyse des performances et des limites", level=1)

doc.add_heading("7.1 Points forts de l'architecture", level=2)
add_bullets([
    ("Latence prédictible",
     "toutes les requêtes principales ciblent une seule partition, ce qui assure des temps de "
     "réponse de l'ordre de la dizaine de millisecondes indépendamment du volume total."),
    ("Débit d'écriture élevé",
     "le BatchStatement par paquet de 100 lignes amortit les allers-retours réseau ; sur un "
     "nœud unique en local, le pipeline soutient sans difficulté plusieurs milliers d'inserts "
     "par seconde."),
    ("Découplage stockage / calcul",
     "les fichiers bruts vivent dans MinIO indépendamment de Cassandra, ce qui autorise un "
     "retraitement complet sans purge de la base."),
    ("Reproductibilité",
     "une seule commande (docker compose up --build) suffit à déployer l'ensemble, ce qui "
     "facilite la validation par un évaluateur et le travail en équipe."),
])

doc.add_heading("7.2 Limites identifiées", level=2)
add_bullets([
    ("Mono-nœud Cassandra",
     "le déploiement actuel utilise un seul nœud avec replication_factor = 1. Cela ne démontre "
     "pas concrètement les propriétés de tolérance aux pannes ; un cluster à trois nœuds serait "
     "nécessaire pour cela."),
    ("Ingestion batch uniquement",
     "le pipeline est conçu comme un job one-shot ; il ne gère pas l'arrivée continue de "
     "nouveaux fichiers. Un consommateur Kafka, un watcher de bucket S3 ou un scheduler "
     "(Airflow) serait à ajouter en production."),
    ("Pas de couche analytique distribuée",
     "le projet ne couvre pas l'étape de transformation lourde (jointures, agrégations "
     "multi-partitions) qui relèverait d'un moteur comme Apache Spark."),
    ("BatchStatement multi-partitions",
     "les batchs actuels sont mono-table mais peuvent toucher plusieurs partitions (plusieurs "
     "pays ou plusieurs marques en un seul batch). Cassandra recommande de limiter chaque batch "
     "à une seule partition pour éviter la pression sur le coordinateur. Un futur correctif "
     "consisterait à grouper les lignes par partition key avant l'envoi."),
    ("Bornes d'ingestion par défaut",
     "les variables CSV_LIMIT, JSON_LIMIT, IMAGE_LIMIT sont fixées à des valeurs de démonstration. "
     "Elles doivent être passées à -1 pour un chargement complet."),
])

doc.add_heading("7.3 Pistes d'évolution", level=2)
add_bullets([
    "Passage à un cluster Cassandra à trois nœuds avec NetworkTopologyStrategy et facteur de réplication 3.",
    "Introduction d'Apache Kafka entre les sources et le service d'ingestion pour passer en quasi-temps réel.",
    "Ajout d'une couche Spark / Trino pour les agrégations cross-partitions à des fins de reporting.",
    "Mise en place de monitoring (Prometheus + Grafana) pour suivre la latence et le débit en production.",
    "Chiffrement at-rest des volumes Docker et chiffrement in-transit (TLS) entre l'ingestor et Cassandra.",
])

add_page_break()

# ===================================================================
# 8. ENJEUX (SÉCURITÉ, RGPD, …)
# ===================================================================
doc.add_heading("8. Discussion des enjeux transverses", level=1)

doc.add_heading("8.1 Sécurité", level=2)
add_para(
    "Dans la configuration actuelle de développement, ni l'authentification Cassandra ni TLS ne "
    "sont activés, et les identifiants par défaut de MinIO (minioadmin / minioadmin) sont "
    "utilisés. Pour un déploiement réel, les actions minimales à appliquer seraient :"
)
add_bullets([
    "activation de l'authentification Cassandra (PasswordAuthenticator + CassandraAuthorizer) et création d'un compte applicatif aux droits restreints au keyspace bigdata_project ;",
    "génération et provisionnement de certificats TLS pour les connexions client→Cassandra et entre nœuds Cassandra ;",
    "rotation des credentials MinIO via le secret manager du fournisseur cloud ou Docker secrets, et désactivation de l'utilisateur root ;",
    "isolation réseau via un réseau Docker dédié et exposition restreinte des ports (suppression des publications 9042 / 9000 si l'accès externe n'est pas nécessaire) ;",
    "journalisation centralisée des accès et alerting sur les tentatives d'authentification échouées.",
])

doc.add_heading("8.2 RGPD et protection des données personnelles", level=2)
add_para(
    "Le projet manipule des données pouvant comporter des informations personnelles "
    "(identifiants client dans la table des ventes, pays, comportement d'achat). Plusieurs "
    "principes du RGPD doivent être pris en compte dans une mise en production :"
)
add_bullets([
    ("Minimisation",
     "ne conserver dans Cassandra que les champs strictement nécessaires aux requêtes "
     "analytiques. Le champ customer_id est aujourd'hui un texte libre ; il pourrait être "
     "remplacé par un identifiant pseudonymisé."),
    ("Droit à l'effacement",
     "Cassandra supporte les DELETE mais ceux-ci génèrent des tombstones qui dégradent les "
     "performances. Une approche plus saine consiste à provisionner un TTL (Time To Live) sur "
     "les lignes contenant des données personnelles, ou à séparer les données identifiantes "
     "dans une table dédiée plus facile à purger."),
    ("Localisation des données",
     "la NetworkTopologyStrategy permet de garantir qu'un facteur de réplication adéquat est "
     "présent dans la zone géographique requise par la réglementation (par exemple, données "
     "UE strictement dans des DC européens)."),
    ("Traçabilité des accès",
     "activer les audit logs Cassandra et les access logs MinIO permet de répondre aux "
     "obligations de redevabilité (accountability)."),
    ("Données brutes dans le Data Lake",
     "les fichiers CSV/JSON archivés sur S3 contiennent les mêmes données personnelles que "
     "Cassandra. Ils doivent être soumis à la même politique de rétention, et chiffrés au "
     "repos (server-side encryption SSE-S3 ou SSE-KMS)."),
])

doc.add_heading("8.3 Scalabilité organisationnelle", level=2)
add_para(
    "Au-delà des aspects purement techniques, l'architecture Big Data choisie facilite le "
    "passage à l'échelle organisationnelle : chaque équipe métier peut développer son propre "
    "modèle de table Cassandra orienté ses requêtes, sans coordination centrale (à la "
    "différence d'un modèle relationnel partagé). Le Data Lake S3 sert de point de référence "
    "commun (single source of truth pour les données brutes) tandis que Cassandra héberge les "
    "vues spécialisées de chaque équipe."
)

doc.add_heading("8.4 Coût et reproductibilité", level=2)
add_para(
    "L'ensemble des briques retenues (Apache Cassandra, MinIO, Docker, Python, librairies "
    "cassandra-driver et boto3) sont sous licences open source. Aucun compte cloud n'est requis "
    "pour exécuter le projet de bout en bout. Cette contrainte du cahier des charges est donc "
    "intégralement respectée, et l'architecture peut être rejouée à coût nul sur tout poste de "
    "travail standard disposant de Docker."
)

add_page_break()

# ===================================================================
# 9. CONCLUSION
# ===================================================================
doc.add_heading("9. Conclusion", level=1)
add_para(
    "Le projet livre une plateforme Big Data fonctionnelle et reproductible répondant à "
    "l'ensemble du périmètre du cahier des charges : ingestion de trois familles de données "
    "(structurées, semi-structurées, non structurées), stockage distribué dans Apache "
    "Cassandra modélisé selon une logique orientée requêtes, archivage dans un Data Lake objet "
    "S3-compatible (MinIO), et orchestration complète via Docker Compose. Les choix techniques "
    "sont systématiquement justifiés par les principes structurants du domaine (5V, théorème "
    "CAP, scalabilité horizontale, séparation stockage / calcul)."
)
add_para(
    "Au-delà du périmètre demandé, l'architecture a été pensée pour évoluer : passage d'un "
    "nœud unique à un cluster multi-DC, remplacement du job batch par un pipeline de "
    "streaming, et intégration future d'une couche analytique distribuée. Le projet constitue "
    "ainsi à la fois une démonstration des concepts vus en cours et un socle réutilisable pour "
    "des cas d'usage industriels plus larges."
)

# ===================================================================
# ANNEXE — DÉPLOIEMENT
# ===================================================================
add_page_break()
doc.add_heading("Annexe A — Instructions de déploiement", level=1)

doc.add_heading("A.1 Prérequis", level=2)
add_bullets([
    "Docker Engine ≥ 24 et Docker Compose v2.",
    "Environ 4 Go de RAM disponibles (Cassandra réclame le plus de mémoire).",
    "5 Go d'espace disque pour les volumes persistants et les images Docker.",
])

doc.add_heading("A.2 Lancement local (MinIO inclus)", level=2)
add_code_block("docker compose up --build")
add_para(
    "Cette commande construit l'image de l'ingestor, démarre Cassandra et MinIO, applique le "
    "schéma CQL, puis exécute le job d'ingestion une fois. Aucun compte AWS n'est requis."
)

doc.add_heading("A.3 Lancement avec AWS S3 réel (optionnel)", level=2)
add_code_block(
    "$env:AWS_ACCESS_KEY_ID=\"<votre_key>\"\n"
    "$env:AWS_SECRET_ACCESS_KEY=\"<votre_secret>\"\n"
    "$env:AWS_REGION=\"eu-west-3\"\n"
    "$env:S3_BUCKET=\"bigdata-raw-amaury\"\n"
    "$env:S3_ENDPOINT_URL=\"\"\n"
    "docker compose up --build"
)

doc.add_heading("A.4 Vérification", level=2)
add_para("Accès à cqlsh pour interroger Cassandra :")
add_code_block(
    "docker exec -it bd_cassandra cqlsh\n"
    "USE bigdata_project;\n"
    "SELECT COUNT(*) FROM sales_by_country_month;"
)
add_para("Console MinIO accessible sur http://localhost:9001 (login minioadmin / minioadmin).")

doc.add_heading("A.5 Structure du dépôt", level=2)
add_code_block(
    ".\n"
    "├── docker-compose.yml\n"
    "├── README.md\n"
    "├── cassandra/\n"
    "│   ├── init.cql           # schéma et keyspace\n"
    "│   └── queries.cql        # requêtes CQL d'exemple\n"
    "├── ingestor/\n"
    "│   ├── Dockerfile\n"
    "│   ├── requirements.txt\n"
    "│   └── ingest.py          # pipeline Python\n"
    "├── data/                  # jeux de données sources\n"
    "└── rapport/\n"
    "    ├── Rapport_Big_Data.docx\n"
    "    └── diagrams/"
)

# ===================================================================
# Sauvegarde
# ===================================================================
doc.save(OUTPUT)
print(f"Rapport généré : {OUTPUT}")
